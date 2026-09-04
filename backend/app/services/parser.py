import io
import re
import uuid
from typing import Dict, Any, List
from pypdf import PdfReader
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from PDF file bytes."""
    text_chunks = []
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from DOCX file bytes."""
    text_chunks = []
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        for p in doc.paragraphs:
            if p.text.strip():
                text_chunks.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_chunks.append(cell.text.strip())
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return "\n".join(text_chunks)


def parse_resume_text_to_data(raw_text: str) -> Dict[str, Any]:
    """
    Parse raw resume text into structured ResumeData dictionary with zero hardcoded fake values.
    """
    lines = [line.strip() for line in raw_text.split("\n") if line.strip()]

    # Extract Contact Info via Regular Expressions
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', raw_text)
    email = email_match.group(0) if email_match else ""

    phone_match = re.search(r'(\+?\d{1,4}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}', raw_text)
    phone = phone_match.group(0) if phone_match else ""

    linkedin_match = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+', raw_text, re.I)
    linkedin = linkedin_match.group(0) if linkedin_match else ""

    github_match = re.search(r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+', raw_text, re.I)
    github = github_match.group(0) if github_match else ""

    portfolio_match = re.search(r'(https?://)?(www\.)?[a-zA-Z0-9_-]+\.(dev|io|me|com|net)', raw_text, re.I)
    portfolio = portfolio_match.group(0) if (portfolio_match and "linkedin" not in portfolio_match.group(0) and "github" not in portfolio_match.group(0)) else ""

    # Full Name heuristic: first line that doesn't contain email or url
    full_name = ""
    for line in lines[:5]:
        if not re.search(r'@|http|www|\.com|\.dev|\d{5}', line) and len(line.split()) <= 4:
            full_name = line
            break

    # Section Headers Detection
    sections: Dict[str, List[str]] = {
        "summary": [],
        "education": [],
        "experience": [],
        "internships": [],
        "projects": [],
        "skills": [],
        "certifications": [],
        "languages": []
    }

    current_section = "summary"
    for line in lines:
        lower_line = line.lower()

        if re.search(r'^(summary|profile|about|career objective)', lower_line):
            current_section = "summary"
            continue
        elif re.search(r'^(education|academic|qualification)', lower_line):
            current_section = "education"
            continue
        elif re.search(r'^(experience|work experience|employment|work history)', lower_line):
            current_section = "experience"
            continue
        elif re.search(r'^(internship|internships)', lower_line):
            current_section = "internships"
            continue
        elif re.search(r'^(projects|project work|key projects)', lower_line):
            current_section = "projects"
            continue
        elif re.search(r'^(skills|technical skills|technologies|skills &)', lower_line):
            current_section = "skills"
            continue
        elif re.search(r'^(certifications|licenses|courses)', lower_line):
            current_section = "certifications"
            continue
        elif re.search(r'^(languages|language proficiency)', lower_line):
            current_section = "languages"
            continue

        sections[current_section].append(line)

    # Process Summary
    summary_text = " ".join(sections["summary"]).strip()
    if summary_text == full_name or summary_text == email:
        summary_text = ""

    # Process Education
    parsed_education = []
    edu_lines = sections["education"]
    if edu_lines:
        current_edu = {
            "id": str(uuid.uuid4()),
            "institution": "",
            "degree": "",
            "fieldOfStudy": "",
            "startDate": "",
            "endDate": "",
            "gpa": "",
            "location": ""
        }
        for l in edu_lines:
            if any(term in l.lower() for term in ["university", "college", "school", "institute", "polytechnic"]):
                if current_edu["institution"]:
                    parsed_education.append(current_edu)
                    current_edu = {
                        "id": str(uuid.uuid4()),
                        "institution": l,
                        "degree": "",
                        "fieldOfStudy": "",
                        "startDate": "",
                        "endDate": "",
                        "gpa": "",
                        "location": ""
                    }
                else:
                    current_edu["institution"] = l
            elif any(deg in l.lower() for deg in ["bachelor", "master", "b.s", "b.tech", "b.e", "m.s", "phd", "degree", "diploma"]):
                current_edu["degree"] = l
            elif "gpa" in l.lower() or "cgpa" in l.lower():
                current_edu["gpa"] = l
            elif not current_edu["institution"]:
                current_edu["institution"] = l
            else:
                if not current_edu["degree"]:
                    current_edu["degree"] = l

        if current_edu["institution"] or current_edu["degree"]:
            parsed_education.append(current_edu)

    # Process Work Experience
    parsed_experience = []
    exp_lines = sections["experience"]
    if exp_lines:
        current_exp = {
            "id": str(uuid.uuid4()),
            "company": "",
            "position": "",
            "location": "",
            "startDate": "",
            "endDate": "",
            "isCurrent": False,
            "bullets": []
        }
        for l in exp_lines:
            if l.startswith("•") or l.startswith("-") or l.startswith("*"):
                bullet = l.lstrip("•-* ").strip()
                if bullet:
                    current_exp["bullets"].append(bullet)
            elif not current_exp["company"]:
                current_exp["company"] = l
            elif not current_exp["position"]:
                current_exp["position"] = l
            else:
                current_exp["bullets"].append(l)

        if current_exp["company"] or current_exp["position"]:
            parsed_experience.append(current_exp)

    # Process Projects
    parsed_projects = []
    proj_lines = sections["projects"]
    if proj_lines:
        current_proj = {
            "id": str(uuid.uuid4()),
            "title": "",
            "description": "",
            "technologies": [],
            "githubLink": "",
            "liveLink": "",
            "bullets": []
        }
        for l in proj_lines:
            if not current_proj["title"]:
                current_proj["title"] = l
            elif l.startswith("•") or l.startswith("-"):
                current_proj["bullets"].append(l.lstrip("•- ").strip())
            elif "tech:" in l.lower() or "technologies:" in l.lower():
                techs = re.sub(r'tech(nologies)?:', '', l, flags=re.I).split(",")
                current_proj["technologies"] = [t.strip() for t in techs if t.strip()]
            else:
                if not current_proj["description"]:
                    current_proj["description"] = l
                else:
                    current_proj["bullets"].append(l)

        if current_proj["title"]:
            parsed_projects.append(current_proj)

    # Process Skills
    parsed_skills = []
    skill_lines = sections["skills"]
    if skill_lines:
        all_skills_flat = []
        for l in skill_lines:
            parts = l.split(":")
            if len(parts) == 2:
                cat_name = parts[0].strip()
                skills_list = [s.strip() for s in parts[1].split(",") if s.strip()]
                parsed_skills.append({
                    "id": str(uuid.uuid4()),
                    "categoryName": cat_name,
                    "skills": skills_list
                })
            else:
                skills_list = [s.strip() for s in l.split(",") if s.strip()]
                all_skills_flat.extend(skills_list)

        if all_skills_flat and not parsed_skills:
            parsed_skills.append({
                "id": str(uuid.uuid4()),
                "categoryName": "Technical Skills",
                "skills": all_skills_flat
            })

    # Assemble parsed ResumeData
    return {
        "personalInfo": {
            "fullName": full_name,
            "email": email,
            "phone": phone,
            "location": "",
            "linkedin": linkedin,
            "github": github,
            "portfolio": portfolio
        },
        "summary": summary_text,
        "education": parsed_education,
        "experience": parsed_experience,
        "internships": [],
        "projects": parsed_projects,
        "skills": parsed_skills,
        "certifications": [],
        "achievements": [],
        "languages": []
    }
